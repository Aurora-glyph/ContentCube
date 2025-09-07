from fastapi import FastAPI, UploadFile, File, HTTPException, status, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
import PyPDF2
from docx import Document
import json
import uuid
import io
from typing import List, Dict, Optional, Union
import re
import csv
from io import StringIO
import os
from dotenv import load_dotenv
import asyncio
import time
import mimetypes
import hashlib
import requests
import tempfile
import subprocess
import base64
from PIL import Image, ImageEnhance
import speech_recognition as sr
import yt_dlp
import ffmpeg
import pytesseract
import whisper
from moviepy import VideoFileClip
import numpy as np
import cv2
from pydub import AudioSegment

# Load environment variables from .env file
load_dotenv()

# Configure Gemini API
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY not found in environment variables. Please check your .env file.")

genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-2.5-flash')

# Initialize Whisper model for audio transcription
try:
    whisper_model = whisper.load_model("base")
except:
    whisper_model = None
    print("Warning: Whisper model not loaded. Audio transcription may be limited.")

app = FastAPI(title="CodeEd Universal Content Repurposer")

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "https://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Enhanced security constants
MAX_FILE_SIZE = 500 * 1024 * 1024  # 500MB for video files
ALLOWED_MIME_TYPES = {
    # Documents
    'application/pdf': ['.pdf'],
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
    'text/plain': ['.txt'],
    
    # Images
    'image/jpeg': ['.jpg', '.jpeg'],
    'image/png': ['.png'],
    'image/webp': ['.webp'],
    'image/gif': ['.gif'],
    'image/bmp': ['.bmp'],
    'image/tiff': ['.tiff', '.tif'],
    
    # Audio
    'audio/mpeg': ['.mp3'],
    'audio/wav': ['.wav'],
    'audio/ogg': ['.ogg'],
    'audio/m4a': ['.m4a'],
    'audio/flac': ['.flac'],
    
    # Video
    'video/mp4': ['.mp4'],
    'video/avi': ['.avi'],
    'video/mov': ['.mov'],
    'video/wmv': ['.wmv'],
    'video/flv': ['.flv'],
    'video/webm': ['.webm'],
    'video/mkv': ['.mkv'],
}

# In-memory storage for demo
jobs_storage = {}

class GeminiRateLimiter:
    def __init__(self, requests_per_minute: int = 10):  # Reduced for multimedia processing
        self.requests_per_minute = requests_per_minute
        self.request_times = []
        
    async def make_request(self, prompt: str, model_instance: any, image_data: bytes = None) -> str:
        """Make a rate-limited request to Gemini API with optional image support"""
        current_time = time.time()
        
        # Remove requests older than 1 minute
        self.request_times = [t for t in self.request_times if current_time - t < 60]
        
        # Wait if we've exceeded the rate limit
        if len(self.request_times) >= self.requests_per_minute:
            sleep_time = 60 - (current_time - self.request_times[0]) + 1
            if sleep_time > 0:
                await asyncio.sleep(sleep_time)
                current_time = time.time()
                self.request_times = [t for t in self.request_times if current_time - t < 60]
        
        try:
            # Make the actual request
            if image_data:
                # For image analysis with Gemini Vision
                image_part = {
                    "mime_type": "image/jpeg",
                    "data": base64.b64encode(image_data).decode()
                }
                response = model_instance.generate_content([prompt, image_part])
            else:
                response = model_instance.generate_content(prompt)
                
            self.request_times.append(current_time)
            return response.text
        except Exception as e:
            # Handle rate limit errors gracefully
            if "rate limit" in str(e).lower():
                await asyncio.sleep(60)  # Wait a minute and retry
                return await self.make_request(prompt, model_instance, image_data)
            raise e

# Global rate limiter instance
rate_limiter = GeminiRateLimiter()

class ContentRequest(BaseModel):
    title: str
    content_type: str = "file"  # "file" or "youtube"
    youtube_url: Optional[str] = None
    language_targets: List[str] = ["hi", "es"]
    style: str = "formal-friendly"

class JobStatus(BaseModel):
    job_id: str
    status: str  # "processing", "completed", "failed"
    progress: int  # 0-100
    result: Optional[Dict] = None
    error: Optional[str] = None

def validate_file(file: UploadFile) -> bool:
    """Enhanced file validation for multimedia files"""
    # Check file size
    if hasattr(file, 'size') and file.size and file.size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds {MAX_FILE_SIZE/1024/1024}MB limit"
        )
    
    # Check file extension
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Filename is required"
        )
    
    file_ext = '.' + file.filename.split('.')[-1].lower()
    allowed_extensions = [ext for exts in ALLOWED_MIME_TYPES.values() for ext in exts]
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type {file_ext} not supported. Supported types: {', '.join(allowed_extensions)}"
        )
    
    return True

def validate_youtube_url(url: str) -> bool:
    """Validate YouTube URL format"""
    youtube_patterns = [
        r'https?://(?:www\.)?youtube\.com/watch\?v=[\w-]+',
        r'https?://youtu\.be/[\w-]+',
        r'https?://(?:www\.)?youtube\.com/embed/[\w-]+',
    ]
    
    return any(re.match(pattern, url) for pattern in youtube_patterns)

def get_file_type(filename: str) -> str:
    """Determine file type category"""
    ext = '.' + filename.split('.')[-1].lower()
    
    for mime_type, extensions in ALLOWED_MIME_TYPES.items():
        if ext in extensions:
            if mime_type.startswith('image/'):
                return 'image'
            elif mime_type.startswith('audio/'):
                return 'audio'
            elif mime_type.startswith('video/'):
                return 'video'
            elif mime_type.startswith('application/') or mime_type.startswith('text/'):
                return 'document'
    
    return 'unknown'

async def extract_text_from_image(image_data: bytes) -> str:
    """Extract text from image using OCR and AI vision"""
    try:
        # Save image temporarily
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            temp_file.write(image_data)
            temp_path = temp_file.name
        
        # OCR extraction
        try:
            ocr_text = pytesseract.image_to_string(Image.open(temp_path))
        except:
            ocr_text = ""
        
        # AI Vision analysis
        vision_prompt = """
        Analyze this image and provide:
        1. A detailed description of what you see
        2. Any text visible in the image
        3. Educational concepts or information that could be extracted
        4. Context that would be useful for learning
        
        Format your response as clear, educational content suitable for creating summaries and quizzes.
        """
        
        try:
            ai_description = await rate_limiter.make_request(vision_prompt, model, image_data)
        except:
            ai_description = "AI vision analysis not available"
        
        # Combine OCR and AI analysis
        combined_text = f"""
        OCR Extracted Text:
        {ocr_text.strip() if ocr_text.strip() else "No text detected"}
        
        Image Analysis:
        {ai_description}
        """
        
        # Clean up temp file
        os.unlink(temp_path)
        
        return combined_text.strip()
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing image: {str(e)}")

async def extract_text_from_audio(audio_data: bytes, filename: str) -> str:
    """Extract text from audio using speech recognition"""
    try:
        # Save audio temporarily
        with tempfile.NamedTemporaryFile(suffix=os.path.splitext(filename)[1], delete=False) as temp_file:
            temp_file.write(audio_data)
            temp_path = temp_file.name
        
        transcript = ""
        
        # Try Whisper first (more accurate)
        if whisper_model:
            try:
                result = whisper_model.transcribe(temp_path)
                transcript = result["text"]
            except Exception as e:
                print(f"Whisper transcription failed: {e}")
        
        # Fallback to speech_recognition
        if not transcript.strip():
            try:
                # Convert to WAV if needed
                audio = AudioSegment.from_file(temp_path)
                wav_path = temp_path.replace(os.path.splitext(temp_path)[1], '.wav')
                audio.export(wav_path, format="wav")
                
                r = sr.Recognizer()
                with sr.AudioFile(wav_path) as source:
                    audio_data = r.record(source)
                    transcript = r.recognize_google(audio_data)
                
                if wav_path != temp_path:
                    os.unlink(wav_path)
                    
            except Exception as e:
                transcript = f"Audio transcription failed: {str(e)}"
        
        # Clean up temp file
        os.unlink(temp_path)
        
        return transcript if transcript.strip() else "No speech detected in audio file"
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing audio: {str(e)}")

async def extract_text_from_video(video_data: bytes, filename: str) -> str:
    """Extract text from video (audio track + key frames)"""
    try:
        # Save video temporarily
        with tempfile.NamedTemporaryFile(suffix=os.path.splitext(filename)[1], delete=False) as temp_file:
            temp_file.write(video_data)
            temp_path = temp_file.name
        
        # Extract audio and transcribe
        try:
            video = VideoFileClip(temp_path)
            audio_path = temp_path.replace(os.path.splitext(temp_path)[1], '.wav')
            
            if video.audio:
                video.audio.write_audiofile(audio_path, verbose=False, logger=None)
                
                # Transcribe audio
                if whisper_model:
                    result = whisper_model.transcribe(audio_path)
                    transcript = result["text"]
                else:
                    transcript = "Audio transcription not available"
                
                os.unlink(audio_path)
            else:
                transcript = "No audio track found in video"
            
            video.close()
            
        except Exception as e:
            transcript = f"Video processing failed: {str(e)}"
        
        # Extract key frames for visual analysis (optional)
        visual_info = ""
        try:
            cap = cv2.VideoCapture(temp_path)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            
            # Extract a few key frames
            key_frames = [0, frame_count // 4, frame_count // 2, 3 * frame_count // 4, frame_count - 1]
            
            for frame_num in key_frames:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_num)
                ret, frame = cap.read()
                if ret:
                    # Convert frame to bytes for analysis
                    _, buffer = cv2.imencode('.jpg', frame)
                    frame_bytes = buffer.tobytes()
                    
                    # Quick AI analysis of frame
                    frame_prompt = "Briefly describe the educational content visible in this video frame."
                    try:
                        frame_analysis = await rate_limiter.make_request(frame_prompt, model, frame_bytes)
                        visual_info += f"Frame {frame_num}: {frame_analysis}\n"
                    except:
                        break  # Stop if rate limited
            
            cap.release()
        except Exception as e:
            visual_info = "Visual analysis not available"
        
        combined_content = f"""
        Video Transcript:
        {transcript}
        
        Visual Content Analysis:
        {visual_info if visual_info.strip() else "Visual analysis not performed"}
        """
        
        # Clean up temp file
        os.unlink(temp_path)
        
        return combined_content.strip()
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing video: {str(e)}")

async def download_youtube_content(url: str) -> str:
    """Download and extract content from YouTube video"""
    try:
        # Configure yt-dlp options
        ydl_opts = {
            'format': 'best[height<=720]',  # Limit quality to manage file size
            'extractaudio': True,
            'audioformat': 'mp3',
            'outtmpl': '/tmp/%(title)s.%(ext)s',
            'quiet': True,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            # Get video info
            info = ydl.extract_info(url, download=False)
            title = info.get('title', 'Unknown')
            description = info.get('description', '')
            duration = info.get('duration', 0)
            
            # Check duration (limit to 30 minutes for processing)
            if duration > 1800:
                raise HTTPException(
                    status_code=400, 
                    detail="Video too long. Please use videos under 30 minutes."
                )
            
            # Download video
            ydl.download([url])
            
            # Find downloaded file
            video_path = None
            for file in os.listdir('/tmp'):
                if file.startswith(title[:20]) and file.endswith(('.mp4', '.webm', '.mkv')):
                    video_path = f'/tmp/{file}'
                    break
            
            if not video_path:
                raise Exception("Downloaded video not found")
            
            # Extract content from downloaded video
            with open(video_path, 'rb') as f:
                video_data = f.read()
            
            content = await extract_text_from_video(video_data, video_path)
            
            # Add metadata
            content = f"""
            Video Title: {title}
            Duration: {duration // 60}:{duration % 60:02d}
            
            Description:
            {description[:500]}...
            
            {content}
            """
            
            # Clean up downloaded file
            os.unlink(video_path)
            
            return content.strip()
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing YouTube video: {str(e)}")

# Keep existing functions for document processing
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file with better handling"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text + "\n"
            except Exception as e:
                print(f"Error extracting page {page_num + 1}: {e}")
                continue
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No readable text found in PDF")
        
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file with better formatting"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            raise HTTPException(status_code=400, detail="No readable text found in DOCX")
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

# Keep existing content generation functions (generate_summary_and_takeaways, generate_mcqs, etc.)
async def generate_summary_and_takeaways(text: str, title: str) -> Dict:
    """Generate summary and key takeaways using Gemini"""
    prompt = f"""
You are an expert instructional designer. Create educational content from this material.

TITLE: {title}

CONTENT: {text[:12000]}

Generate a JSON response with this exact structure:
{{
    "summary": "A 250-300 word summary suitable for undergraduate students. Make it engaging, clear, and well-structured.",
    "takeaways": ["5 key bullet points", "each 10-15 words", "covering main concepts", "in simple language", "actionable insights"]
}}

Keep the summary factual, concise, and scaffolded for learning. Make takeaways specific and memorable.
Generate the summary in different paragraphs and bullets, etc.
"""
    
    try:
        response_text = await rate_limiter.make_request(prompt, model)
        content = response_text.strip()
        content = re.sub(r'```json\n?', '', content)
        content = re.sub(r'```\n?', '', content)
        return json.loads(content)
    except Exception as e:
        return {
            "summary": f"Summary generation failed: {str(e)}",
            "takeaways": ["Error in processing", "Please try again", "Check input format", "Ensure content quality", "Contact support if needed"]
        }

async def generate_mcqs(text: str, title: str) -> List[Dict]:
    """Generate MCQs using Gemini"""
    prompt = f"""
You are an expert quiz creator. Create 10 high-quality multiple-choice questions from this educational content.

TITLE: {title}
CONTENT: {text[:12000]}

Generate a JSON array with this exact structure:
[
    {{
        "question": "Clear, specific question text",
        "options": ["A) First option", "B) Second option", "C) Third option", "D) Fourth option"],
        "correct_answer": "B",
        "explanation": "Brief 1-2 sentence explanation of why this is correct",
        "bloom_level": "Remember|Understand|Apply|Analyze|Evaluate|Create"
    }}
]

Rules:
- Cover different sections of the content
- One clearly correct answer per question
- Plausible distractors (wrong options)
- No negative phrasing ("Which is NOT...")
- Explanations under 25 words
- Mix of bloom taxonomy levels
"""
    
    try:
        response_text = await rate_limiter.make_request(prompt, model)
        content = response_text.strip()
        content = re.sub(r'```json\n?', '', content)
        content = re.sub(r'```\n?', '', content)
        return json.loads(content)
    except Exception as e:
        return [{
            "question": f"Question generation failed: {str(e)}",
            "options": ["A) Error", "B) Please try again", "C) Check content", "D) Contact support"],
            "correct_answer": "A",
            "explanation": "System error occurred",
            "bloom_level": "Remember"
        }]

async def generate_flashcards(text: str, title: str) -> List[Dict]:
    """Generate flashcards using Gemini"""
    prompt = f"""
Create 6 flashcards for spaced repetition study from this educational content.

TITLE: {title}
CONTENT: {text[:12000]}

Generate a JSON array with this exact structure:
[
    {{
        "front": "Question or concept prompt (clear and specific)",
        "back": "Answer or explanation (under 30 words, self-contained)"
    }}
]

Rules:
- Each flashcard should test one atomic concept
- Front side: question, term, or scenario
- Back side: concise answer or definition
- No cross-references between cards
- Cover key concepts from the content
"""
    
    try:
        response_text = await rate_limiter.make_request(prompt, model)
        content = response_text.strip()
        content = re.sub(r'```json\n?', '', content)
        content = re.sub(r'```\n?', '', content)
        return json.loads(content)
    except Exception as e:
        return [{
            "front": f"Flashcard generation failed: {str(e)}",
            "back": "Please try again with different content"
        }]

async def localize_content(content: Dict, target_language: str) -> Dict:
    """Translate content to target language"""
    lang_names = {"hi": "Hindi", "es": "Spanish", "fr": "French"}
    lang_name = lang_names.get(target_language, target_language)
    
    prompt = f"""
Translate and culturally adapt this educational content to {lang_name} for local learners.

CONTENT: {json.dumps(content, indent=2)}

Generate the same JSON structure but with all text translated to {lang_name}.

Rules:
- Keep technical terms accurate
- Adapt cultural references appropriately  
- Maintain academic tone
- Keep same JSON structure
- Simplify complex idioms
- Ensure translations are natural for learners
"""
    
    try:
        response_text = await rate_limiter.make_request(prompt, model)
        content_text = response_text.strip()
        content_text = re.sub(r'```json\n?', '', content_text)
        content_text = re.sub(r'```\n?', '', content_text)
        return json.loads(content_text)
    except Exception as e:
        return {**content, "translation_error": f"Translation failed: {str(e)}"}

def create_google_forms_csv(mcqs: List[Dict]) -> str:
    """Create CSV format for Google Forms import"""
    output = StringIO()
    writer = csv.writer(output)
    
    writer.writerow(["Question", "Option 1", "Option 2", "Option 3", "Option 4", "Correct Answer", "Explanation"])
    
    for mcq in mcqs:
        if "options" in mcq and len(mcq["options"]) >= 4:
            options_text = [opt[3:] if len(opt) > 3 else opt for opt in mcq["options"]]
            writer.writerow([
                mcq["question"],
                options_text[0],
                options_text[1], 
                options_text[2],
                options_text[3],
                mcq["correct_answer"],
                mcq["explanation"]
            ])
    
    return output.getvalue()

@app.post("/repurpose")
async def create_repurpose_job(
    file: Optional[UploadFile] = File(None),
    title: str = Form("Educational Content"),
    content_type: str = Form("file"),
    youtube_url: Optional[str] = Form(None)
):
    """Enhanced content repurposing endpoint supporting multiple file types and YouTube"""
    
    # Validate input
    if content_type == "youtube":
        if not youtube_url or not validate_youtube_url(youtube_url):
            raise HTTPException(status_code=400, detail="Valid YouTube URL required")
    else:
        if not file:
            raise HTTPException(status_code=400, detail="File required when content_type is 'file'")
        validate_file(file)
    
    # Generate job ID
    job_id = str(uuid.uuid4())
    
    # Initialize job status
    jobs_storage[job_id] = {
        "job_id": job_id,
        "status": "processing",
        "progress": 0,
        "result": None,
        "error": None
    }
    
    try:
        # Extract content based on type
        if content_type == "youtube":
            jobs_storage[job_id]["progress"] = 10
            text = await download_youtube_content(youtube_url)
            content_source = "YouTube"
        else:
            file_content = await file.read()
            file_type = get_file_type(file.filename)
            
            jobs_storage[job_id]["progress"] = 10
            
            if file_type == "document":
                if file.filename.lower().endswith('.pdf'):
                    text = extract_text_from_pdf(file_content)
                elif file.filename.lower().endswith('.docx'):
                    text = extract_text_from_docx(file_content)
                else:  # .txt
                    text = file_content.decode('utf-8')
            elif file_type == "image":
                text = await extract_text_from_image(file_content)
            elif file_type == "audio":
                text = await extract_text_from_audio(file_content, file.filename)
            elif file_type == "video":
                text = await extract_text_from_video(file_content, file.filename)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type")
            
            content_source = file_type.title()
        
        jobs_storage[job_id]["progress"] = 30
        
        # Generate educational content
        summary_data = await generate_summary_and_takeaways(text, title)
        jobs_storage[job_id]["progress"] = 50
        
        mcqs = await generate_mcqs(text, title)
        jobs_storage[job_id]["progress"] = 70
        
        flashcards = await generate_flashcards(text, title)
        jobs_storage[job_id]["progress"] = 85
        
        # Create result
        result = {
            "title": title,
            "content_source": content_source,
            "summary": summary_data.get("summary", ""),
            "takeaways": summary_data.get("takeaways", []),
            "mcqs": mcqs,
            "flashcards": flashcards,
            "localized": {},
            "exports": {},
        }
        
        # Generate localizations (simplified for demo due to rate limits)
        try:
            for lang in ["hi", "es"]:
                localized = await localize_content({
                    "summary": result["summary"],
                    "takeaways": result["takeaways"][:2],  # Limit to reduce API calls
                }, lang)
                result["localized"][lang] = {
                    **localized,
                    "mcqs": result["mcqs"],
                    "flashcards": result["flashcards"]
                }
        except:
            # Provide placeholder if translation fails
            result["localized"] = {
                "hi": {"summary": "हिंदी अनुवाद सेवा अस्थायी रूप से सीमित है।", "takeaways": ["अनुवाद सेवा सीमित"]},
                "es": {"summary": "Servicio de traducción temporalmente limitado.", "takeaways": ["Servicio limitado"]}
            }
        
        # Create exports
        result["exports"]["google_forms_csv"] = create_google_forms_csv(mcqs)
        result["exports"]["full_json"] = json.dumps(result, indent=2)
        
        # Complete job
        jobs_storage[job_id].update({
            "status": "completed",
            "progress": 100,
            "result": result
        })
        
    except Exception as e:
        jobs_storage[job_id].update({
            "status": "failed",
            "progress": 0,
            "error": str(e)
        })
    
    return {"job_id": job_id}

# Keep existing endpoints
@app.get("/jobs/{job_id}")
async def get_job_status(job_id: str):
    """Get job status and progress"""
    if job_id not in jobs_storage:
        raise HTTPException(status_code=404, detail="Job not found")
    return jobs_storage[job_id]

@app.get("/outputs/{job_id}")
async def get_job_outputs(job_id: str):
    """Get completed job outputs"""
    if job_id not in jobs_storage:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs_storage[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Job not completed yet")
    
    return job["result"]

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "message": "CodeEd Universal Content Repurposer API is running"}

@app.get("/")
async def root():
    return {"message": "CodeEd Universal Content Repurposer API", "status": "running", "supported_types": list(ALLOWED_MIME_TYPES.keys())}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)